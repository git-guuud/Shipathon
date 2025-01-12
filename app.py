import streamlit as st
import os
import aspose.words as aw
import shutil
import google.generativeai as genai
import fitz 
from PIL import Image
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from dotenv import load_dotenv
import traceback

load_dotenv(".env")
API_KEY = os.getenv("API")

def getjson(pdf_file):
    pdf_document = fitz.open(pdf_file)
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel("gemini-2.0-flash-exp")

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        image.save(f"uploads/page_{page_num}.jpg", "JPEG")
        response = model.generate_content(["Act like a text scanner and translator. Extract text as it is without analyzing it and without summarizing it. Treat all images as a whole document and analyze them accordingly. Think of it as a document with multiple pages, each image being a page. Understand page-to-page flow logically and semantically. Translate the given text to English if its not in english and return it.", image], stream=True, generation_config={"temperature": 0.01})
        response.resolve()
        schema = '''
            {
            "document": [
                {
                    "page_number": 1,
                    "elements": [
                    {
                        "text": "",
                        "font_size": 12,
                        "alignment": "",
                        "font_style": ""
                    },
                    {
                        "text": "",
                        "font_size": 12,
                        "alignment": "",
                        "font_style": ""
                    }
                    ]
                }
                ]
            }
            '''
        response = model.generate_content(["Act like a text formatter. Analyse the image and understand the formatting of the text. Understand the flow of the document, check where a paragraph ends don't make every line a paragraph. Apply the formatting on the translated version of the text. Return a JSON file with given schema that gives the translated text formatted like the original image.", image, response.text, schema], stream=True, generation_config={"temperature": 0.01})
        response.resolve()
        file = open(f"uploads/page_{page_num}.json", 'w', encoding="utf-8")
        file.write(response.text.split("json")[1].split('`')[0].strip())
        file.close()
    return len(pdf_document)

def apply_paragraph_style(paragraph, font_size=None, alignment=None, font_style=None):
    if paragraph.runs[0]:
        if font_size:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
    
        if font_style:
            if "bold" in font_style.lower():
                paragraph.runs[0].bold = True
            if "italic" in font_style.lower():
                paragraph.runs[0].italic = True
        
        if alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT
            }
            paragraph.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

def create_docx_from_json(json_file_path, output_docx_path):
    with open(json_file_path, "r", encoding="utf-8") as file:
        data = json.load(file)
    
    document = Document(output_docx_path)

    for page in data.get("document", []):
        for element in page.get("elements", []):
            text = element.get("text", "")
            font_size = element.get("font_size", None)
            alignment = element.get("alignment", None)
            font_style = element.get("font_style", None)
            
            paragraph = document.add_paragraph(text)
            apply_paragraph_style(paragraph, font_size=font_size, alignment=alignment, font_style=font_style)
    
    document.save(output_docx_path)

def createdoc(name, length):
    output_docx = 'uploads/' + name 
    doc = Document()
    doc.save(f"uploads/{name}")
    for i in range(length):
        injson = 'uploads/page_' + str(i) + ".json"
        input_json = injson  # Replace with your uploaded JSON file path
        create_docx_from_json(input_json, output_docx)


def main():
    try:
        st.title("PDF To DOCX Translator")
        if os.path.exists("uploads"):
            shutil.rmtree("uploads")
        os.makedirs("uploads")
        pdf_file = None
        pdf_file = st.file_uploader("Please upload a PDF file to translate.", type=["pdf"])

        if pdf_file is not None:
            with open(os.path.join("uploads",pdf_file.name),"wb") as f: 
                f.write(pdf_file.getbuffer())  

            with st.spinner('Uploading and Translating'):
                length = getjson(f"uploads/{pdf_file.name}")
                docx_file = pdf_file.name.rsplit(".",1)[0] + ".docx"
                createdoc(docx_file ,length)
                doc = aw.Document(f"uploads/{docx_file}")
                save_options = aw.saving.ImageSaveOptions(aw.SaveFormat.JPEG)
                save_options.page_set = aw.saving.PageSet(0)
                image_file = pdf_file.name.rsplit(".",1)[0] +  ".jpg"
                doc.save(f"uploads/{image_file}", save_options)
                st.subheader("Preview")
                st.image(f"uploads/{image_file}")
                with open(f"uploads/{docx_file}", 'rb') as f:
                    document = f.read()
                st.download_button('Download Translated Docx', data = document, file_name=docx_file)
            st.success("Done!")
    except Exception as error:
        st.error("Something went wrong. Please try again later.")
        print(traceback.format_exc())

main()